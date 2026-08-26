"""Pure document parsing, chunking, and search-text preparation helpers."""

from hashlib import sha256
import io
from re import sub
from typing import Any, cast


def sha256_hex(data: bytes) -> str:
    """Return the stable content checksum used for upload deduplication."""
    return sha256(data).hexdigest()


def decode_document_text(raw: bytes, content_type: str | None, filename: str) -> str:
    """Extract plain text from a supported uploaded document."""
    lower_name = filename.lower()
    if lower_name.endswith(".pdf") or content_type == "application/pdf":
        text = extract_pdf_text(raw)
    elif lower_name.endswith(".docx") or content_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        if lower_name.endswith(".doc"):
            raise ValueError(
                "Legacy .doc is not supported. Please convert to .docx before uploading."
            )
        text = extract_docx_text(raw)
    elif lower_name.endswith(".xlsx") or content_type in (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ):
        if lower_name.endswith(".xls"):
            raise ValueError(
                "Legacy .xls is not supported. Please convert to .xlsx before uploading."
            )
        text = extract_xlsx_text(raw)
    else:
        supported_by_name = lower_name.endswith(
            (".txt", ".md", ".markdown", ".csv", ".json", ".log")
        )
        supported_by_type = bool(
            content_type
            and (
                content_type.startswith("text/")
                or content_type
                in {
                    "application/json",
                    "application/x-ndjson",
                    "application/csv",
                }
            )
        )
        if not supported_by_name and not supported_by_type:
            raise ValueError(
                "Unsupported file type. Supported: txt, md, csv, json, log, pdf, docx, xlsx."
            )
        text = raw.decode("utf-8", errors="replace").strip()
    if not text:
        raise ValueError("Document is empty after text decoding.")
    return text


def extract_pdf_text(raw: bytes) -> str:
    """Extract PDF text through pypdf, PyMuPDF, then Tesseract OCR."""
    for extractor in (extract_pdf_text_pypdf, extract_pdf_text_fitz, extract_pdf_text_ocr):
        text = extractor(raw)
        if text:
            return text
    raise ValueError(
        "PDF text extraction returned no text. The PDF may be a scanned image "
        "document and OCR is unavailable or failed. Ensure tesseract-ocr and "
        "tesseract-ocr-chi-sim are installed."
    )


def extract_pdf_text_pypdf(raw: bytes) -> str:
    """Tier 1: pypdf text extraction."""
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception:
        return ""
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n\n".join(page.strip() for page in pages if page.strip()).strip()


def extract_pdf_text_fitz(raw: bytes) -> str:
    """Tier 2: PyMuPDF text extraction for complex layouts."""
    try:
        import fitz
    except ImportError:
        return ""
    try:
        document = fitz.open(stream=raw, filetype="pdf")
    except Exception:
        return ""
    pages = []
    for page in document:
        page_text = (page.get_text("text") or "").strip()
        if page_text:
            pages.append(page_text)
    document.close()
    return "\n\n".join(pages).strip()


def extract_pdf_text_ocr(raw: bytes) -> str:
    """Tier 3: render PDF pages with PyMuPDF and OCR with Tesseract."""
    try:
        import fitz
        import pytesseract
        from PIL import Image
    except ImportError:
        return ""
    try:
        document = fitz.open(stream=raw, filetype="pdf")
    except Exception:
        return ""
    pages = []
    for page in document:
        pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
        image = Image.frombytes("RGB", cast(Any, (pix.width, pix.height)), pix.samples)
        page_text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
        if page_text:
            pages.append(page_text)
    document.close()
    return "\n\n".join(pages).strip()


def extract_docx_text(raw: bytes) -> str:
    """Extract paragraphs and table rows from a Word .docx."""
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "python-docx is required to parse Word documents but is not installed."
        ) from exc
    try:
        document = Document(io.BytesIO(raw))
    except Exception as exc:
        raise ValueError(f"Failed to read .docx: {exc}") from exc
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def extract_xlsx_text(raw: bytes) -> str:
    """Extract sheet names and nonempty rows from an Excel .xlsx."""
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise RuntimeError(
            "openpyxl is required to parse Excel documents but is not installed."
        ) from exc
    try:
        workbook = load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    except Exception as exc:
        raise ValueError(f"Failed to read .xlsx: {exc}") from exc
    parts: list[str] = []
    for sheet in workbook.worksheets:
        sheet_title = str(sheet.title or "").strip()
        if sheet_title:
            parts.append(f"# {sheet_title}")
        for row in sheet.iter_rows(values_only=True):
            cells = [str(cell).strip() for cell in row if cell is not None and str(cell).strip()]
            if cells:
                parts.append(" | ".join(cells))
    workbook.close()
    return "\n\n".join(parts).strip()


def chunk_text(text: str, *, chunk_size: int, overlap: int) -> list[str]:
    """Split normalized text into bounded, overlap-aware chunks."""
    safe_chunk_size = min(max(chunk_size, 200), 4000)
    safe_overlap = min(max(overlap, 0), safe_chunk_size // 2)
    normalized = "\n".join(line.rstrip() for line in text.splitlines()).strip()
    paragraphs = [part.strip() for part in normalized.split("\n\n") if part.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs or [normalized]:
        if len(paragraph) > safe_chunk_size:
            if current:
                chunks.append(current.strip())
                current = ""
            start = 0
            while start < len(paragraph):
                end = min(len(paragraph), start + safe_chunk_size)
                chunks.append(paragraph[start:end].strip())
                if end == len(paragraph):
                    break
                start = max(end - safe_overlap, start + 1)
            continue
        candidate = f"{current}\n\n{paragraph}".strip() if current else paragraph
        if len(candidate) <= safe_chunk_size:
            current = candidate
        else:
            chunks.append(current.strip())
            prefix = current[-safe_overlap:].strip() if safe_overlap and current else ""
            current = f"{prefix}\n\n{paragraph}".strip() if prefix else paragraph
    if current:
        chunks.append(current.strip())
    return [chunk for chunk in chunks if chunk]


def normalize_search_text(value: str) -> str:
    return sub(r"\s+", " ", sub(r"[^\w\u4e00-\u9fff]+", " ", value.lower())).strip()


def query_terms(query: str) -> list[str]:
    """Produce bounded Western-word and CJK-bigram terms for PostgreSQL FTS."""
    normalized = normalize_search_text(query)
    terms = [term for term in normalized.split(" ") if len(term) >= 2]
    cjk_chars = [char for char in normalized if "\u4e00" <= char <= "\u9fff"]
    cjk_bigrams = [f"{left}{right}" for left, right in zip(cjk_chars, cjk_chars[1:])]
    result: list[str] = []
    for term in [*terms, *cjk_bigrams]:
        if term not in result:
            result.append(term)
    return result[:24]


def score_chunk(search_text: str, query_terms_: list[str], raw_query: str) -> float:
    normalized_query = normalize_search_text(raw_query)
    if not query_terms_:
        return 0.1
    score = 0.0
    if normalized_query and normalized_query in search_text:
        score += 0.45
    for term in query_terms_:
        occurrences = search_text.count(term)
        if occurrences:
            score += min(0.2, 0.08 * occurrences)
    return round(min(score, 1.0), 4)


def build_fts_text(value: str) -> str:
    """Build FTS text from CJK bigrams and Western terms."""
    return " ".join(query_terms(value))


def build_fts_query(query: str) -> str:
    """Build a recall-oriented OR tsquery from the normalized query terms."""
    return " | ".join(query_terms(query))


def rough_token_count(text: str) -> int:
    return max(1, len(text) // 4)
